"""Renderers turn structured content into real office files.

Openability is the contract; byte stability is explicitly not.
"""

from pathlib import Path

from workbench.artifacts.render import RenderOutcome, render_document
from workbench.core.artifacts import (
    FormattedDocument,
    HeadingBlock,
    ListBlock,
    ParagraphBlock,
    SpreadsheetContent,
    SpreadsheetSheet,
    TableBlock,
)


def sheet_content() -> str:
    return SpreadsheetContent(
        sheets=(
            SpreadsheetSheet(
                name="Fees",
                columns=("Matter", "Hours"),
                rows=(("Vantage NDA", 3.5), ("Harbor lease", 2)),
            ),
        )
    ).canonical_json()


def formatted_content() -> str:
    return FormattedDocument(
        blocks=(
            HeadingBlock(level=1, text="Engagement Letter"),
            ParagraphBlock(text="Scope of work below."),
            ListBlock(ordered=False, items=("Review", "Redline")),
            TableBlock(columns=("Item", "Amount"), rows=(("Retainer", "$5,000"),)),
        )
    ).canonical_json()


def test_markdown_passes_through(tmp_path: Path) -> None:
    outcome = render_document("markdown", "# Hello", tmp_path / "note.md")
    assert outcome == RenderOutcome(path=tmp_path / "note.md", skipped=None)
    assert (tmp_path / "note.md").read_text(encoding="utf-8") == "# Hello"


def test_spreadsheet_renders_openable_xlsx(tmp_path: Path) -> None:
    from openpyxl import load_workbook

    outcome = render_document("spreadsheet", sheet_content(), tmp_path / "fees.xlsx")
    assert outcome.skipped is None
    workbook = load_workbook(outcome.path)
    sheet = workbook["Fees"]
    assert [cell.value for cell in sheet[1]] == ["Matter", "Hours"]
    assert [cell.value for cell in sheet[2]] == ["Vantage NDA", 3.5]
    assert [cell.value for cell in sheet[3]] == ["Harbor lease", 2]


def test_formatted_renders_openable_docx(tmp_path: Path) -> None:
    from docx import Document

    outcome = render_document(
        "formatted", formatted_content(), tmp_path / "letter.docx"
    )
    assert outcome.skipped is None
    document = Document(str(outcome.path))
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert texts[0] == "Engagement Letter"
    assert "Scope of work below." in texts
    table = document.tables[0]
    assert table.rows[0].cells[0].text == "Item"
    assert table.rows[1].cells[1].text == "$5,000"


def test_pdf_renders_without_any_external_converter(tmp_path: Path) -> None:
    """A .pdf must be a PDF everywhere, not a .docx wearing the name.

    This used to shell out to LibreOffice and leave a docx behind when
    `soffice` was missing — which is the normal case on a developer
    machine, in CI, and inside the task container, so the firm never
    produced a single PDF.
    """

    outcome = render_document("formatted", formatted_content(), tmp_path / "letter.pdf")

    assert outcome.path == tmp_path / "letter.pdf"
    assert outcome.skipped is None
    assert outcome.path.read_bytes().startswith(b"%PDF-")


def test_sheet_titles_excel_would_reject_are_repaired(tmp_path: Path) -> None:
    """An author naming a tab "Revenue/Expenses" is being reasonable.

    openpyxl raises on the slash, and that exception took the whole
    materialization of a world down with it.
    """

    content = SpreadsheetContent(
        sheets=(
            SpreadsheetSheet(name="Revenue/Expenses", columns=("A",), rows=(("1",),)),
            SpreadsheetSheet(name="Revenue/Expenses", columns=("A",), rows=(("2",),)),
            SpreadsheetSheet(name="X" * 40, columns=("A",), rows=(("3",),)),
        )
    ).canonical_json()

    outcome = render_document("spreadsheet", content, tmp_path / "wp.xlsx")

    from openpyxl import load_workbook

    titles = load_workbook(outcome.path).sheetnames
    assert titles == ["Revenue-Expenses", "Revenue-Expenses-2", "X" * 31]
