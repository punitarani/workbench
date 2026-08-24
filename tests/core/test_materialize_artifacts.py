"""Materialization renders structured documents as real office files."""

from pathlib import Path

from core.artifacts import (
    FormattedDocument,
    HeadingBlock,
    ParagraphBlock,
    SpreadsheetContent,
    SpreadsheetSheet,
)
from core.events import Event
from core.events.control import SimRunStartedPayload
from core.events.documents import DocumentCreatedPayload
from core.events.people import PersonRecordPayload
from core.worldlog import WorldLogWriter
from environment import materialize

FEES = SpreadsheetContent(
    sheets=(
        SpreadsheetSheet(
            name="Fees",
            columns=("Matter", "Hours"),
            rows=(("Vantage NDA", 3.5),),
        ),
    )
).canonical_json()

LETTER = FormattedDocument(
    blocks=(
        HeadingBlock(level=1, text="Engagement Letter"),
        ParagraphBlock(text="Scope of work."),
    )
).canonical_json()


def _document(
    document_id: str, path: str, content_format: str, content: str
) -> DocumentCreatedPayload:
    return DocumentCreatedPayload(
        kind="document.created",
        document_id=document_id,
        author="per-ann-liu",
        title=path.rsplit("/", 1)[-1],
        path=path,
        location="repository",
        content_format=content_format,
        content=content,
    )


def artifact_events(*, with_pdf: bool = False) -> list[Event]:
    payloads = [
        SimRunStartedPayload(
            kind="sim.run.started",
            run_id="run-artifacts",
            seed_root=7,
            workplace_id="mini",
            config_hash="0" * 64,
            schema_version=1,
            epoch="2026-03-12T00:00:00+00:00",
            timezone="UTC",
        ),
        PersonRecordPayload(
            kind="person.record",
            person_id="per-ann-liu",
            name="Ann Liu",
            email_address="ann@mini.example",
            title="Counsel",
            department="Legal",
            manager=None,
            affiliation="internal",
            timezone="UTC",
        ),
        _document("doc-000001", "/legal/billing/fees.xlsx", "spreadsheet", FEES),
        _document("doc-000002", "/legal/letters/engagement.docx", "formatted", LETTER),
    ]
    if with_pdf:
        payloads.append(
            _document("doc-000003", "/legal/letters/terms.pdf", "formatted", LETTER)
        )
    return [
        Event(seq=seq, time=0, tag=payload.kind, source="gm", payload=payload)
        for seq, payload in enumerate(payloads)
    ]


def write_log(tmp_path: Path, events: list[Event]) -> Path:
    log_path = tmp_path / "world.jsonl"
    with WorldLogWriter(log_path) as writer:
        for event in events:
            writer.append(event)
    return log_path


def test_materialize_renders_real_office_files(tmp_path: Path) -> None:
    from docx import Document
    from openpyxl import load_workbook

    out = tmp_path / "bundle"
    result = materialize(write_log(tmp_path, artifact_events()), out)

    workbook = load_workbook(out / "workspace" / "legal" / "billing" / "fees.xlsx")
    assert [cell.value for cell in workbook["Fees"][2]] == ["Vantage NDA", 3.5]

    letters = out / "workspace" / "legal" / "letters"
    document = Document(str(letters / "engagement.docx"))
    texts = [p.text for p in document.paragraphs if p.text]
    assert texts[0] == "Engagement Letter"

    assert result.document_files == 2
    assert result.skipped_renders == ()


def test_a_pdf_document_materializes_as_a_pdf(tmp_path: Path) -> None:
    out = tmp_path / "bundle"
    result = materialize(write_log(tmp_path, artifact_events(with_pdf=True)), out)

    rendered = out / "workspace" / "legal" / "letters" / "terms.pdf"
    assert rendered.exists()
    assert rendered.read_bytes().startswith(b"%PDF-")
    assert result.skipped_renders == ()
